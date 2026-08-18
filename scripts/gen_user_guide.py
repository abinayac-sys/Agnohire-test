# -*- coding: utf-8 -*-
"""Generates AgnoHire_User_Guide.docx — a corporate-grade, editable Word user
guide with cover page, document control, page-numbered footers, running headers,
branded headings, callout boxes, styled tables, and figure placeholders.

Run:  python scripts/gen_user_guide.py
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- palette ---------------------------------------------------------------
INDIGO = RGBColor.from_string("4F46E5")
INDIGO_LT = "EEF0FE"
INK = RGBColor.from_string("1F2937")
MUTED = RGBColor.from_string("6B7280")
WHITE = RGBColor.from_string("FFFFFF")
LINE = "D8DCE8"
BOX_FILL = "F7F9FC"
BOX_LINE = "C7CDDA"

doc = Document()
FIG = [0]  # figure counter

# ---- base styles -----------------------------------------------------------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

doc.styles['Title'].font.name = 'Calibri'
for h, size, color in [('Heading 1', 18, INDIGO), ('Heading 2', 13.5, INDIGO),
                       ('Heading 3', 11.5, INK)]:
    st = doc.styles[h]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(12)
    st.paragraph_format.space_after = Pt(4)


# ---- low-level helpers -----------------------------------------------------
def _shade(tc, fill):
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    pr.append(shd)


def _cell_borders(cell, color=BOX_LINE, sz=4, left=None, left_sz=None):
    tcPr = cell._tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement('w:' + edge)
        c = left if (edge == 'left' and left) else color
        s = left_sz if (edge == 'left' and left_sz) else sz
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(s))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), c)
        b.append(e)
    tcPr.append(b)


def _row_height(row, twips):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement('w:trHeight'); h.set(qn('w:val'), str(twips)); h.set(qn('w:hRule'), 'atLeast')
    trPr.append(h)


def _p_border(paragraph, color="4F46E5", sz=6, space=6):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), str(space)); bottom.set(qn('w:color'), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _field(run, instr):
    a = OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'), 'begin')
    b = OxmlElement('w:instrText'); b.set(qn('xml:space'), 'preserve'); b.text = instr
    c = OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'), 'separate')
    d = OxmlElement('w:fldChar'); d.set(qn('w:fldCharType'), 'end')
    for el in (a, b, c, d):
        run._r.append(el)


# ---- content helpers -------------------------------------------------------
def h1(t):
    p = doc.add_heading(t, level=1)
    _p_border(p, color="4F46E5", sz=6, space=6)
    return p


def h2(t): return doc.add_heading(t, level=2)
def h3(t): return doc.add_heading(t, level=3)
def para(t): return doc.add_paragraph(t)


def steps(items):
    for it in items:
        doc.add_paragraph(it, style='List Number')


def bullets(items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(it, tuple):
            r = p.add_run(it[0] + " — "); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def callout(kind, text):
    cfg = {'Note': (INDIGO_LT, "4F46E5"), 'Tip': ("E7F6EC", "047857"),
           'Important': ("FEF3E2", "B45309"), 'Caution': ("FDECEC", "B91C1C")}
    fill, accent = cfg[kind]
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0); cell.width = Inches(6.5)
    _shade(cell._tc, fill)
    _cell_borders(cell, color=fill, sz=4, left=accent, left_sz=24)
    p = cell.paragraphs[0]
    r = p.add_run(kind.upper() + "   "); r.bold = True
    r.font.color.rgb = RGBColor.from_string(accent); r.font.size = Pt(9)
    rt = p.add_run(text); rt.font.size = Pt(10)
    doc.add_paragraph()


def screenshot(caption, height=2200):
    FIG[0] += 1
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _row_height(t.rows[0], height)
    cell = t.cell(0, 0); cell.width = Inches(6.5)
    _shade(cell._tc, BOX_FILL)
    _cell_borders(cell, color=BOX_LINE, sz=8)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[  Insert screenshot here  ]")
    r.italic = True; r.font.color.rgb = RGBColor.from_string("9AA1B2"); r.font.size = Pt(10.5)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run("Figure %d.  %s" % (FIG[0], caption))
    cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = MUTED
    doc.add_paragraph()


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(headers):
        c = t.cell(0, i); _shade(c._tc, "4F46E5"); _cell_borders(c, color="4F46E5", sz=4)
        p = c.paragraphs[0]; r = p.add_run(htext); r.bold = True
        r.font.color.rgb = WHITE; r.font.size = Pt(9.5)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        fill = "FFFFFF" if ri % 2 == 0 else "F3F5FB"
        for i, val in enumerate(row):
            _shade(cells[i]._tc, fill); _cell_borders(cells[i], color="E4E8F3", sz=4)
            p = cells[i].paragraphs[0]; r = p.add_run(str(val)); r.font.size = Pt(9.5)
            if i == 0:
                r.bold = True
    if widths:
        for i, w in enumerate(widths):
            for cell in t.columns[i].cells:
                cell.width = Inches(w)
    doc.add_paragraph()
    return t


def page_break():
    doc.add_page_break()


# ===========================================================================
# PAGE SETUP — margins, header & footer with page numbers
# ===========================================================================
sec = doc.sections[0]
sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(1.0); sec.right_margin = Inches(1.0)
sec.different_first_page_header_footer = True  # cover has no header/footer

# Running header (right-aligned title, thin rule)
hp = sec.header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("AgnoHire  ·  User Guide")
hr.font.size = Pt(8); hr.font.color.rgb = MUTED
_p_border(hp, color=LINE, sz=4, space=2)

# Footer: "Confidential   |   Page X of Y"
fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
f1 = fp.add_run("Confidential — © 2026 AgnoHire     |     Page ")
f1.font.size = Pt(8); f1.font.color.rgb = MUTED
pr = fp.add_run(); pr.font.size = Pt(8); pr.font.color.rgb = MUTED; _field(pr, "PAGE")
f2 = fp.add_run(" of "); f2.font.size = Pt(8); f2.font.color.rgb = MUTED
nr = fp.add_run(); nr.font.size = Pt(8); nr.font.color.rgb = MUTED; _field(nr, "NUMPAGES")

# ===========================================================================
# COVER PAGE
# ===========================================================================
cl = doc.add_paragraph(); cl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
clr = cl.add_run("CONFIDENTIAL"); clr.bold = True; clr.font.size = Pt(9)
clr.font.color.rgb = RGBColor.from_string("B91C1C")

for _ in range(3):
    doc.add_paragraph()

# logo placeholder (compact, centered)
lt = doc.add_table(rows=1, cols=1); lt.alignment = WD_TABLE_ALIGNMENT.CENTER
_row_height(lt.rows[0], 1500)
lc = lt.cell(0, 0); lc.width = Inches(2.4)
_shade(lc._tc, BOX_FILL); _cell_borders(lc, color=BOX_LINE, sz=8)
lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
lpr = lp.add_run("[ Company logo ]"); lpr.italic = True
lpr.font.color.rgb = RGBColor.from_string("9AA1B2"); lpr.font.size = Pt(10)

doc.add_paragraph()
t1 = doc.add_paragraph(); t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
t1r = t1.add_run("AgnoHire"); t1r.font.size = Pt(46); t1r.bold = True; t1r.font.color.rgb = INDIGO
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2r = t2.add_run("User Guide"); t2r.font.size = Pt(24); t2r.font.color.rgb = INK
t3 = doc.add_paragraph(); t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
t3r = t3.add_run("Enterprise AI Recruitment Platform")
t3r.italic = True; t3r.font.size = Pt(12); t3r.font.color.rgb = MUTED

for _ in range(6):
    doc.add_paragraph()

# cover info strip
ct = doc.add_table(rows=1, cols=3); ct.alignment = WD_TABLE_ALIGNMENT.CENTER
covinfo = [("VERSION", "1.0"), ("STATUS", "Released"), ("DATE", "June 2026")]
for i, (k, v) in enumerate(covinfo):
    c = ct.cell(0, i); _shade(c._tc, "4F46E5"); _cell_borders(c, color="4F46E5", sz=4)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rk = p.add_run(k + "\n"); rk.font.size = Pt(8); rk.font.color.rgb = RGBColor.from_string("C7D2FE")
    rv = p.add_run(v); rv.bold = True; rv.font.size = Pt(12); rv.font.color.rgb = WHITE
    c.width = Inches(2.1)
page_break()

# ===========================================================================
# DOCUMENT CONTROL
# ===========================================================================
h1("Document Control")
h2("Document Information")
table(["Field", "Detail"], [
    ("Document title", "AgnoHire — User Guide"),
    ("Product", "AgnoHire — Enterprise AI Recruitment Platform"),
    ("Version", "1.0"),
    ("Status", "Released"),
    ("Classification", "Confidential"),
    ("Owner", "[ Document owner / team ]"),
    ("Prepared by", "[ Author name ]"),
    ("Approved by", "[ Approver name ]"),
    ("Date", "June 2026"),
], widths=[2.0, 4.5])

h2("Revision History")
table(["Version", "Date", "Author", "Summary of changes"], [
    ("0.1", "[ date ]", "[ author ]", "Initial draft."),
    ("1.0", "June 2026", "[ author ]", "First released edition."),
    ("", "", "", ""),
], widths=[1.0, 1.3, 1.6, 2.6])

h2("Confidentiality & Legal Notice")
para("This document and the information it contains are confidential and proprietary to AgnoHire. "
     "It is provided solely for the use of authorised personnel and may not be copied, distributed, "
     "or disclosed to any third party, in whole or in part, without prior written permission.")
para("All product names, screens, and workflows described herein are subject to change as the "
     "platform evolves. While care has been taken to ensure accuracy at the time of writing, this "
     "guide is provided for informational purposes and without warranty of any kind.")
callout("Note", "Replace the bracketed [ … ] placeholders above (owner, author, approver, dates) "
                "with your organisation's details before circulating this document.")
page_break()

# ===========================================================================
# TABLE OF CONTENTS
# ===========================================================================
h1("Table of Contents")
para("")
tp = doc.add_paragraph()
run = tp.add_run()
a = OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'), 'begin')
b = OxmlElement('w:instrText'); b.set(qn('xml:space'), 'preserve'); b.text = 'TOC \\o "1-2" \\h \\z \\u'
c = OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'), 'separate')
tt = OxmlElement('w:t')
tt.text = "Right-click anywhere in this line and choose “Update Field” to generate the contents."
d = OxmlElement('w:fldChar'); d.set(qn('w:fldCharType'), 'end')
for el in (a, b, c, tt, d):
    run._r.append(el)
page_break()

# ===========================================================================
# PREFACE
# ===========================================================================
h1("Preface")
h2("About This Guide")
para("This guide explains how to use AgnoHire to manage the complete recruitment lifecycle — from "
     "creating a job opening through sourcing, screening, interviewing, panel review, offers, and "
     "onboarding. It is written for everyday users and focuses on the tasks you perform in the "
     "application rather than on technical implementation.")

h2("Intended Audience")
bullets([
    ("Recruiters & sourcers", "creating jobs, screening candidates, and running interviews."),
    ("HR & hiring managers", "reviewing pipelines, panels, offers, onboarding, and compliance."),
    ("Administrators", "managing users, roles, sectors, branding, and system settings."),
    ("Candidates", "using the candidate portal to apply, interview, and get help."),
])
para("No prior experience with AgnoHire is assumed. Familiarity with using a modern web browser is "
     "sufficient.")

h2("How This Guide Is Organised")
para("The guide moves in roughly the order of the hiring lifecycle. Each chapter opens with a short "
     "introduction, then provides task-based, step-by-step instructions. Reference material — a "
     "roles-and-permissions matrix, a glossary, FAQs, and troubleshooting — is collected in the "
     "appendices.")

h2("Conventions Used in This Guide")
para("The following callout boxes highlight information of particular importance:")
callout("Note", "Additional context or a useful detail.")
callout("Tip", "A shortcut or recommended way of working.")
callout("Important", "Something you should not overlook to complete a task correctly.")
callout("Caution", "An action that may have consequences — proceed carefully.")
para("Interface labels such as buttons, menus, and field names are written in the order you click "
     "them, for example: Admin → System Config → General.")

h2("Working with the Screenshots")
para("Every grey box labelled “[ Insert screenshot here ]” is a placeholder with a numbered caption. "
     "To add an image, click inside the box, delete the placeholder text, then use Insert → Pictures. "
     "Update the caption text to describe the screenshot. You may add, remove, reorder, or rebrand "
     "any part of this document to suit your organisation.")
page_break()

# ===========================================================================
# 1. INTRODUCTION & OVERVIEW
# ===========================================================================
h1("1. Introduction & Overview")
para("AgnoHire is a complete recruitment platform that manages the entire hiring journey in one "
     "place. Artificial-intelligence features assist throughout — drafting job descriptions, scoring "
     "résumés, summarising interviews, and answering questions — but the platform remains fully "
     "functional when AI is switched off.")
h2("Key Capabilities")
bullets([
    ("Jobs & approvals", "create requisitions, route them for approval, and generate descriptions with AI."),
    ("Screening", "upload résumés, auto-score fit, and shortlist applicants."),
    ("Sourcing", "referrals, channels, and curated candidate lists."),
    ("Interviewing", "AI interviews, live interviews, panel reviews, scheduling, and reminders."),
    ("Pipeline", "a drag-and-drop board that tracks every candidate's stage."),
    ("Offers & onboarding", "draft, send, and track offers, then manage onboarding and checks."),
    ("Assistant", "Agno, a built-in chatbot for staff and candidates."),
    ("Compliance", "audit logs and data-protection (GDPR) tooling."),
    ("Administration", "users, roles, sectors, branding, and system configuration."),
])
callout("Note", "AgnoHire is organised by “sectors” (business units). Most users see only their own "
                "sector's data; administrators can work across all sectors.")

# ===========================================================================
# 2. CORE CONCEPTS
# ===========================================================================
h1("2. Core Concepts")
para("A few concepts appear throughout AgnoHire. Understanding them makes the rest of the guide "
     "easier to follow.")
bullets([
    ("Sector", "a business unit and the boundary for data separation. Candidates, jobs, and reports "
     "belong to a sector."),
    ("Role", "a named set of permissions assigned to a user (for example, Recruiter or HR)."),
    ("Permission", "a single capability, such as viewing candidates or managing offers."),
    ("Pipeline stage", "where a candidate sits in the hiring process, shown as a board column."),
    ("Graceful AI", "if no AI key is configured, AI suggestions are skipped and the platform still "
     "works normally."),
])
callout("Tip", "If a menu item or button you expect is missing, it is almost always because your "
               "role does not include the matching permission. Ask an administrator for access.")

# ===========================================================================
# 3. GETTING STARTED
# ===========================================================================
h1("3. Getting Started")
h2("3.1 Before You Begin")
bullets([
    ("Browser", "a current version of Chrome, Edge, Firefox, or Safari."),
    ("Account", "an email and password (or single sign-on) provided by your administrator."),
    ("Access", "your role determines which areas you can see."),
])
callout("Tip", "Quick start: sign in → open the area you need from the sidebar → use the primary "
               "button on each page (for example “New Job” or “Add Candidate”) to create records.")

h2("3.2 Signing In")
steps([
    "Open the AgnoHire web address in your browser.",
    "Enter your email address and password, or choose “Continue with Google” if enabled.",
    "Select “Sign in”. You arrive at the home page for your role.",
])
callout("Note", "If you forget your password, an administrator can reset it from the Admin Console.")
screenshot("The sign-in screen")

h2("3.3 Finding Your Way Around")
para("After signing in you will see three main areas:")
bullets([
    ("Sidebar (left)", "the navigation menu; the links shown depend on your role."),
    ("Top bar", "the page title area, the notification bell, and your account menu."),
    ("Main area", "the content of the current page."),
])
screenshot("The main interface — sidebar, top bar, and content area")

h2("3.4 Notifications")
para("The bell icon in the top bar shows alerts such as new assignments, panel invitations, and "
     "offer responses. A badge indicates unread items; select the bell to read them.")

h2("3.5 Signing Out")
para("Open your account menu in the top-right corner and choose “Sign out”. For shared computers, "
     "always sign out when you finish.")

# ===========================================================================
# 4. DASHBOARD & ANALYTICS
# ===========================================================================
h1("4. Dashboard & Analytics")
para("The analytics area summarises hiring performance: headline numbers, the hiring funnel, and "
     "trends over time, with options to filter and export.")
steps([
    "Open Analytics from the sidebar.",
    "Review the summary cards (for example open jobs, candidates, interviews, and offers).",
    "Examine the hiring funnel and time-series charts for trends.",
    "Apply filters to focus on a date range or category.",
    "Select Export to download a report as CSV, or save a snapshot for later comparison.",
])
screenshot("The analytics dashboard")

# ===========================================================================
# 5. JOBS & REQUISITIONS
# ===========================================================================
h1("5. Jobs & Requisitions")
para("A job (requisition) represents an opening you are hiring for. Jobs can pass through an "
     "approval workflow before they are published.")
h2("5.1 Create a Job")
steps([
    "Open Jobs from the sidebar and select “New Job”.",
    "Complete the title, sector and domain, location, work mode, and other details.",
    "Optionally select “Generate with AI” to draft the description, then refine it.",
    "Save as a draft, or submit for approval.",
])
screenshot("Creating a new job")
h2("5.2 Approve or Reject a Job")
steps([
    "Open a job that is pending approval.",
    "Review its details and description.",
    "Select Approve to publish, or Reject with a reason to return it to the author.",
])
callout("Tip", "Set up reusable templates and approver lists so new jobs start from a consistent base.")

# ===========================================================================
# 6. CANDIDATES & SCREENING
# ===========================================================================
h1("6. Candidates & Screening")
para("Candidates are the people applying to, or being considered for, your jobs.")
h2("6.1 Add a Candidate")
steps([
    "Open Candidates from the sidebar and select “Add Candidate”.",
    "Enter their name, contact details, and the related job.",
    "Upload their résumé. AgnoHire parses it and can automatically score how well they fit the role.",
    "Save. The candidate now appears in your list and pipeline.",
])
screenshot("Adding a candidate and uploading a résumé")
h2("6.2 Search & Shortlist")
steps([
    "Use the search box and filters to find candidates by name, skills, or status.",
    "Open a candidate to view their profile, résumé, and fit score.",
    "Assign them to a recruiter or advance them along the pipeline.",
])
screenshot("The candidate list with search and filters")

# ===========================================================================
# 7. SOURCING
# ===========================================================================
h1("7. Sourcing")
para("Sourcing brings candidates in through referrals, channels, and curated lists.")
bullets([
    ("Referrals", "record who referred a candidate and track the referral's progress."),
    ("Channels", "track where candidates originate (job boards, agencies, and so on)."),
    ("Lists", "build curated shortlists and assign them in bulk."),
])
screenshot("The sourcing area")

# ===========================================================================
# 8. THE HIRING PIPELINE
# ===========================================================================
h1("8. The Hiring Pipeline")
para("The pipeline is a visual board showing where each candidate sits. Columns represent stages "
     "such as Sourced, Applied, Screening, Interview, Offer, Hired, and Rejected.")
steps([
    "Open Pipeline from the sidebar and choose a job.",
    "Drag a candidate card between columns to change their stage.",
    "Select a card to open the candidate and add notes.",
    "Mark a note as private so only you can see it.",
])
callout("Note", "The board updates live: if a colleague moves a candidate, you see it without refreshing.")
screenshot("The drag-and-drop pipeline board")

# ===========================================================================
# 9. INTERVIEWS
# ===========================================================================
h1("9. Interviews")
h2("9.1 AI Interviews")
para("AI interviews let candidates answer questions through a secure link, with automatic grading "
     "for multiple-choice questions and AI scoring for written or coding answers.")
steps([
    "Open Interviews and create an interview from a question bank.",
    "Share the secure link with the candidate.",
    "When they finish, review their answers and scores.",
])
screenshot("Reviewing an AI interview")
h2("9.2 Scheduling Interviews")
steps([
    "Open Scheduling and choose an available slot; working hours and buffers are respected.",
    "Confirm the interview. Reminders are sent automatically before it begins.",
])
screenshot("Scheduling an interview")

# ===========================================================================
# 10. SKILL ASSESSMENTS
# ===========================================================================
h1("10. Skill Assessments")
para("Assessments are tests you assign to candidates to measure specific skills.")
steps([
    "Open Assessments and build a test from your question banks.",
    "Assign it to one or more candidates.",
    "Candidates take the test through a secure link; scoring is automatic, with AI help for open answers.",
    "Review results against your pass mark.",
])
screenshot("Building and assigning an assessment")

# ===========================================================================
# 11. HIRING PANEL & FEEDBACK
# ===========================================================================
h1("11. Hiring Panel & Feedback")
para("For panel interviews, several reviewers provide structured feedback that is combined into a "
     "single recommendation.")
steps([
    "Open Panel Reviews and select the interview.",
    "Add panel members if required; each receives an invitation to accept or decline.",
    "Each reviewer submits star-rating feedback and a recommendation.",
    "The system displays the combined (weighted) consensus.",
    "The decision-maker records the final outcome.",
])
screenshot("Submitting panel feedback")

# ===========================================================================
# 12. VIDEO INTERVIEW INTELLIGENCE
# ===========================================================================
h1("12. Video Interview Intelligence")
para("When an interview transcript is available, AgnoHire produces useful metrics — such as speaking "
     "time and key topics — and, with AI enabled, a summary and communication scores. A reviewer can "
     "add notes and a recommendation.")
screenshot("An interview intelligence report")

# ===========================================================================
# 13. OFFERS & ONBOARDING
# ===========================================================================
h1("13. Offers & Onboarding")
h2("13.1 Make an Offer")
steps([
    "Open Offers and select “Create Offer”, then choose the candidate's application.",
    "Complete the offer details and attach the offer letter.",
    "Save as a draft, then select Send when ready.",
])
screenshot("Creating an offer")
h2("13.2 Track the Response & Onboarding")
steps([
    "When a candidate accepts, their application automatically moves to Hired and an onboarding "
    "checklist is created.",
    "Open the Onboarding tab to track background verification and complete checklist items.",
])
callout("Important", "Only draft offers can be edited. Once an offer is sent, edit it by withdrawing "
                     "and reissuing, per your organisation's process.")
screenshot("Onboarding checklist and background verification")

# ===========================================================================
# 14. THE AI ASSISTANT (AGNO)
# ===========================================================================
h1("14. The AI Assistant (Agno)")
para("Agno is a built-in assistant that answers common questions. It first looks for a matching "
     "answer in your FAQ knowledge base, then — if AI is enabled — generates a helpful reply.")
h2("14.1 For Staff")
bullets([
    ("Try Agno", "a demo chat to test answers."),
    ("FAQ admin", "add and edit the questions and answers Agno uses."),
    ("Transcripts", "review past conversations within your sector."),
])
screenshot("The Agno assistant and FAQ administration")
h2("14.2 For Candidates")
para("Candidates reach Agno from the Support page in their portal to ask about applications, "
     "interviews, and the hiring process.")
screenshot("The candidate support chat")

# ===========================================================================
# 15. COMPLIANCE
# ===========================================================================
h1("15. Compliance (Audit & Data Protection)")
h2("15.1 Audit Logs")
steps([
    "Open Audit Logs.",
    "Filter by action, type, person, or date, and search the text.",
    "Open an entry to see exactly what changed.",
    "Export the results to CSV when required.",
])
screenshot("The audit log viewer")
h2("15.2 Data-Subject (GDPR) Requests")
bullets([
    ("Access / Portability", "produce a portable data bundle for a candidate."),
    ("Deletion", "permanently erase a candidate's personal data."),
    ("Consent & Retention", "record consent and define data-retention rules."),
])
callout("Caution", "Deletion permanently and irreversibly erases a candidate's personal data. "
                   "Confirm the request is valid before proceeding.")
screenshot("The compliance and data-protection area")

# ===========================================================================
# 16. ADMINISTRATION
# ===========================================================================
h1("16. Administration")
para("The Admin Console is where administrators configure the platform. Each area requires the "
     "matching permission.")
h2("16.1 Users")
steps([
    "Open Admin → Users.",
    "Select “New User”, enter details, and assign a role and sector.",
    "Use the row actions to edit, deactivate, or reset a password.",
])
screenshot("User management")
h2("16.2 Roles & Permissions")
para("Open Admin → Roles to review each role and select the permissions it should hold. The Super "
     "Admin role always holds every permission and cannot be edited.")
screenshot("The role and permission matrix")
h2("16.3 Sectors & Domains")
para("Create and manage business sectors and their domains under Admin → Sectors.")
screenshot("Sectors and domains")
h2("16.4 Integrations")
para("Add third-party connections under Admin → Integrations. Secret keys are stored encrypted and "
     "shown masked.")
screenshot("Integrations")
h2("16.5 Email Templates")
para("Edit the wording of automated emails under Admin → Email Templates. Starter templates are "
     "provided.")
screenshot("Email templates")
h2("16.6 System Configuration")
steps([
    "Open Admin → System Config.",
    "Settings are grouped by category (General, Security, Email, AI, and more).",
    "Change a value and select the save icon on that row.",
    "Use “Send test” to verify your email (SMTP) settings.",
])
callout("Note", "Settings are stored in the database and take effect at runtime — no redeployment is "
                "required. Secret values such as API keys are encrypted and shown masked.")
screenshot("System configuration")

h2("16.7 Branding — App Icon & Company Logo")
para("You can brand AgnoHire with your own images from Admin → System Config → General. Two separate "
     "image settings are available:")
bullets([
    ("App icon", "the square mark shown next to the name in the sidebar and on the sign-in screen. "
     "It fills the square, so a roughly square image (for example 256 × 256) works best."),
    ("Company logo", "a wider logo shown as the main image on the sign-in screen."),
])
steps([
    "Open Admin → System Config and locate the General category.",
    "Next to “App icon”, select Upload and choose a PNG, JPG, or WEBP image.",
    "The badge in the sidebar updates immediately.",
    "Next to “Company logo”, select Upload to set the wider sign-in logo.",
    "To return to the default (your company initial), select the remove icon.",
])
screenshot("Uploading an app icon and company logo")
screenshot("The branded sidebar and sign-in screen")

# ===========================================================================
# 17. CANDIDATE PORTAL
# ===========================================================================
h1("17. The Candidate Portal")
para("Candidates have a simplified area where they can view their applications, take interviews and "
     "assessments through secure links, and chat with the Agno assistant for help.")
screenshot("The candidate portal")

# ===========================================================================
# 18. USING AGNOHIRE ON MOBILE
# ===========================================================================
h1("18. Using AgnoHire on a Mobile Device")
para("AgnoHire adapts to small screens. On a phone or narrow window, the navigation menu is hidden "
     "to maximise space and opens on demand.")
steps([
    "Tap the menu icon (☰) in the top-left corner to open the navigation drawer.",
    "Select a destination; the drawer closes automatically.",
    "Tap outside the drawer, or the close icon, to dismiss it without navigating.",
    "Wide tables scroll sideways — swipe left or right within a table to see more columns.",
])
callout("Tip", "All features are available on mobile. For data-heavy tasks such as building "
               "assessments, a larger screen is more comfortable.")
screenshot("AgnoHire on a mobile device — the navigation drawer")

# ===========================================================================
# APPENDIX A — ROLES & PERMISSIONS
# ===========================================================================
page_break()
h1("Appendix A — Roles & Permissions Reference")
para("What you can see and do depends on your role. The principal roles are summarised below.")
table(["Role", "Typically can", "Sees"], [
    ("Super Admin", "Everything; manages the platform itself.", "All sectors"),
    ("Admin", "Manage users, roles, sectors, integrations, settings; view audit logs.", "All sectors"),
    ("HR", "Manage candidates, offers, onboarding, and compliance.", "Own sector"),
    ("Recruiter", "Manage jobs, candidates, sourcing, interviews, and the pipeline.", "Own sector"),
    ("Hiring Manager", "Review the pipeline; give panel feedback; record decisions.", "Own sector"),
    ("Panel Member", "Provide interview feedback for assigned panels.", "Own sector"),
    ("Candidate", "Use the portal; take interviews and assessments; ask the assistant.", "Self only"),
], widths=[1.4, 3.6, 1.5])

# ===========================================================================
# APPENDIX B — GLOSSARY
# ===========================================================================
h1("Appendix B — Glossary")
table(["Term", "Meaning"], [
    ("Sector", "A business unit; the boundary for data separation."),
    ("Requisition", "A job opening you are hiring for."),
    ("Pipeline", "The stages a candidate moves through, shown as a board."),
    ("Panel", "A group of reviewers who assess a candidate together."),
    ("Consensus", "A combined recommendation derived from panel feedback."),
    ("Offer", "A formal job offer sent to a candidate."),
    ("Onboarding", "The steps after a candidate accepts an offer."),
    ("BGV", "Background verification, part of onboarding."),
    ("GDPR", "Data-protection rules; AgnoHire includes tools to honour data requests."),
    ("Audit log", "A record of who did what, and when."),
    ("FAQ", "The question-and-answer knowledge base used by the Agno assistant."),
], widths=[1.6, 4.9])

# ===========================================================================
# APPENDIX C — FAQ
# ===========================================================================
h1("Appendix C — Frequently Asked Questions")
faq = [
    ("I can't see a menu item that a colleague has.",
     "Menu items depend on your role and permissions. Contact an administrator if you need access."),
    ("I can only see some candidates or jobs.",
     "Most users are limited to their own sector. Administrators see all sectors."),
    ("The AI features don't appear to do anything.",
     "AI is optional. If no AI key is configured, the platform still works but skips AI suggestions. "
     "An administrator can add a key under System Config → AI."),
    ("My logo looks cropped.",
     "The app icon fills a square, so very wide or tall images are cropped. Use a roughly square "
     "image for the app icon, and the Company logo setting for wide logos."),
    ("I forgot my password.",
     "Ask an administrator to reset it from Admin → Users."),
]
for q, ans in faq:
    p = doc.add_paragraph(); r = p.add_run("Q.  " + q); r.bold = True
    doc.add_paragraph("A.  " + ans)

# ===========================================================================
# APPENDIX D — TROUBLESHOOTING
# ===========================================================================
h1("Appendix D — Troubleshooting")
table(["Symptom", "Likely cause", "What to do"], [
    ("A page won't load.", "Temporary connection issue.", "Refresh the page; if it persists, contact support."),
    ("“Unauthorised” after clicking a link.", "Your role lacks the permission.", "Request access from an administrator."),
    ("No email was received.", "Email settings not configured.", "Ask an administrator to verify SMTP under System Config."),
    ("A résumé won't upload.", "Unsupported file type or size.", "Use PDF, DOCX, or TXT within the size limit."),
], widths=[1.9, 1.9, 2.7])

# ===========================================================================
# APPENDIX E — SUPPORT & CONTACT
# ===========================================================================
h1("Appendix E — Support & Contact")
para("For help that this guide does not cover, contact your internal support team:")
table(["Channel", "Detail"], [
    ("Support email", "[ support@yourcompany.com ]"),
    ("Help desk / portal", "[ link ]"),
    ("Administrator", "[ name / team ]"),
    ("Hours", "[ business hours ]"),
], widths=[2.0, 4.5])
doc.add_paragraph()
end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = end.add_run("— End of User Guide —"); er.italic = True; er.font.color.rgb = MUTED

# ---- save ------------------------------------------------------------------
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "AgnoHire_User_Guide.docx")
doc.save(out)
print("Saved:", out, "| figures:", FIG[0])
